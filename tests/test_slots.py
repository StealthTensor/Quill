"""Tests for blank detection and answer write-back."""

from __future__ import annotations

import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from worksheetfiller import identity, writer  # noqa: E402
from worksheetfiller.config import Student  # noqa: E402
from worksheetfiller.llm import LLMError, parse_json_object  # noqa: E402
from worksheetfiller.slots import (  # noqa: E402
    KIND_CELL,
    KIND_PARA,
    KIND_STUB,
    extract,
    is_stub,
    strip_trailing_slots,
)


@pytest.fixture
def student() -> Student:
    return Student(
        name="Asha Rao",
        reg_no="23A91A0501",
        branch="CSE",
        section="B",
        persona="hostel student",
    )


def build_worksheet() -> Document:
    """A miniature worksheet with each kind of blank."""
    document = Document()

    table = document.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "Session"
    table.cell(0, 1).text = "1"
    table.cell(0, 2).text = "Lecture"
    table.cell(0, 3).text = "Lecture 1"
    table.cell(1, 0).text = "Name"
    table.cell(1, 2).text = "Reg. No."

    document.add_paragraph("C1. Five key take-aways:")
    document.add_paragraph("1.")
    document.add_paragraph("2.")

    document.add_paragraph("Q1. Explain the difference between animals and human beings.")
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("")

    grid = document.add_table(rows=2, cols=3)
    grid.cell(0, 0).text = ""
    grid.cell(0, 1).text = "Do I WANT this?"
    grid.cell(0, 2).text = "Is this my STATE now?"
    grid.cell(1, 0).text = "To be happy"

    return document


class TestExtraction:
    def test_finds_every_kind_of_blank(self):
        extraction = extract(build_worksheet())
        kinds = [slot.kind for slot in extraction.slots]

        assert kinds.count(KIND_STUB) == 2
        assert KIND_PARA in kinds
        assert kinds.count(KIND_CELL) >= 4

    def test_consecutive_blank_lines_become_one_slot(self):
        extraction = extract(build_worksheet())
        para_slots = [slot for slot in extraction.slots if slot.kind == KIND_PARA]

        assert len(para_slots) == 1
        assert len(para_slots[0].paragraphs) == 3

    def test_markers_appear_in_reading_order(self):
        extraction = extract(build_worksheet())
        text = extraction.text
        positions = [text.index(f"<<{slot.id}>>") for slot in extraction.slots]

        assert positions == sorted(positions)
        assert "Q1. Explain the difference" in text

    def test_cell_slot_carries_row_and_column_labels(self):
        extraction = extract(build_worksheet())
        labelled = {
            (slot.label, slot.column) for slot in extraction.slots if slot.kind == KIND_CELL
        }

        assert ("Name", "1") in labelled
        assert ("To be happy", "Do I WANT this?") in labelled

    def test_blank_matrix_corner_is_not_a_slot(self):
        extraction = extract(build_worksheet())
        corners = [
            slot
            for slot in extraction.slots
            if slot.kind == KIND_CELL and not slot.label and not slot.column
        ]

        assert corners == []

    def test_blank_in_a_header_row_without_labels_is_still_a_slot(self):
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(1, 0).text = "Date"

        extraction = extract(document)
        assert len(extraction.slots) == 2

    def test_trailing_padding_is_dropped(self):
        document = Document()
        document.add_paragraph("Q1. Something?")
        document.add_paragraph("")

        extraction = extract(document)
        assert len(extraction.slots) == 1
        assert strip_trailing_slots(extraction) == ()

    @pytest.mark.parametrize("text", ["1.", "2)", "-", "•", "  3.  "])
    def test_stub_patterns(self, text):
        assert is_stub(text)

    @pytest.mark.parametrize("text", ["1. Something", "Q1.", "hello"])
    def test_non_stub_patterns(self, text):
        assert not is_stub(text)


class TestIdentity:
    def test_identity_cells_are_resolved_locally(self, student):
        extraction = extract(build_worksheet())
        known, pending = identity.resolve(extraction.slots, student, "12-08-2026")

        assert "Asha Rao" in known.values()
        assert "23A91A0501" in known.values()
        assert all(slot.id not in known for slot in pending)

    @pytest.mark.parametrize(
        "label,expected",
        [
            ("Reg. No.", "reg_no"),
            ("Roll Number", "reg_no"),
            ("Branch / Sec.", "branch"),
            ("Date", "date"),
            ("Why?", None),
        ],
    )
    def test_label_aliases(self, label, expected):
        assert identity.field_for_label(label) == expected


class TestWriter:
    def test_answers_land_in_the_document(self, student):
        document = build_worksheet()
        extraction = extract(document)
        answers = {slot.id: f"answer {slot.id}" for slot in extraction.slots}

        written = writer.apply_answers(extraction.slots, answers)
        assert written == len(extraction.slots)

        body = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "answer" in body
        assert document.tables[0].cell(1, 1).text.strip() != ""

    def test_multiline_answer_expands_into_extra_paragraphs(self):
        document = Document()
        document.add_paragraph("Q1. Explain.")
        document.add_paragraph("")
        document.add_paragraph("Next heading")

        extraction = extract(document)
        slot = extraction.slots[0]
        writer.apply_answers((slot,), {slot.id: "line one\nline two\nline three"})

        texts = [paragraph.text for paragraph in document.paragraphs]
        assert "line one" in texts and "line two" in texts and "line three" in texts
        assert texts.index("line three") < texts.index("Next heading")

    def test_stub_answer_keeps_the_number(self):
        document = Document()
        document.add_paragraph("1.")

        extraction = extract(document)
        writer.apply_answers(extraction.slots, {extraction.slots[0].id: "Values matter."})

        assert document.paragraphs[0].text.strip() == "1. Values matter."

    def test_empty_answers_are_skipped(self):
        document = Document()
        document.add_paragraph("1.")

        extraction = extract(document)
        assert writer.apply_answers(extraction.slots, {extraction.slots[0].id: "   "}) == 0


class TestJsonParsing:
    def test_plain_object(self):
        assert parse_json_object('{"S1": "a"}') == {"S1": "a"}

    def test_fenced_object(self):
        assert parse_json_object('```json\n{"S1": "a"}\n```') == {"S1": "a"}

    def test_object_with_surrounding_prose(self):
        reply = 'Here you go:\n{"S1": "a", "S2": "b"}\nHope that helps.'
        assert parse_json_object(reply) == {"S1": "a", "S2": "b"}

    def test_braces_inside_strings_do_not_break_parsing(self):
        assert parse_json_object('{"S1": "a } b"}') == {"S1": "a } b"}

    def test_missing_object_raises(self):
        with pytest.raises(LLMError):
            parse_json_object("no json here")
