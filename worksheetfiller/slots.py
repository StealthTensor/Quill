"""Find the blanks in a .docx worksheet and render it as prompt-ready text.

The document is walked in reading order. Every blank we can write into becomes
a Slot with a stable id. The same walk produces a plain-text rendering of the
document where each blank is replaced by its `<<S3>>` marker, so the model sees
every question in its original context.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from typing import Iterator

from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

KIND_CELL = "cell"
KIND_PARA = "para"
KIND_STUB = "stub"

MARKER_TEMPLATE = "<<{slot_id}>>"
MARKER_PATTERN = re.compile(r"<<(S\d+)>>")

# "1." / "2)" / "-" / "•" on a line of its own: a numbered blank waiting for text.
STUB_PATTERN = re.compile(r"^\s*(?:\d+\s*[.)]|[-*•●])\s*$")


@dataclass(frozen=True)
class Slot:
    """One writable blank in the document."""

    id: str
    kind: str
    paragraphs: tuple[Paragraph, ...]
    label: str = ""
    column: str = ""
    max_chars: int | None = None

    def hint(self) -> str:
        parts = [self.kind]
        if self.label:
            parts.append(f'row "{self.label}"')
        if self.column:
            parts.append(f'column "{self.column}"')
        if self.max_chars:
            parts.append(f"max {self.max_chars} chars")
        return ", ".join(parts)


@dataclass
class Extraction:
    """Result of scanning a document."""

    text: str
    slots: tuple[Slot, ...]

    @property
    def slot_ids(self) -> tuple[str, ...]:
        return tuple(slot.id for slot in self.slots)

    def by_id(self) -> dict[str, Slot]:
        return {slot.id: slot for slot in self.slots}


class _SlotNamer:
    def __init__(self) -> None:
        self._count = 0

    def next(self) -> str:
        self._count += 1
        return f"S{self._count}"


def iter_block_items(parent: DocumentObject | _Cell) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables of `parent` in document order."""
    if isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        parent_element = parent.element.body

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def cell_text(cell: _Cell) -> str:
    return " ".join(part.strip() for part in cell.text.split("\n")).strip()


def is_blank(text: str) -> bool:
    return not text or not text.strip()


def is_stub(text: str) -> bool:
    return bool(STUB_PATTERN.match(text))


def _unique_cells(row) -> list[_Cell]:
    """Row cells with horizontally merged duplicates removed."""
    seen: set[int] = set()
    cells: list[_Cell] = []
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        cells.append(cell)
    return cells


def _writable_paragraph(cell: _Cell) -> Paragraph:
    """The paragraph inside a cell that an answer should be written into."""
    if cell.paragraphs:
        return cell.paragraphs[-1]
    return cell.add_paragraph()


def _has_corner_stub(grid: list[list[_Cell]], header: list[str]) -> bool:
    """True when cell (0, 0) is the blank corner of a labelled matrix.

    Such a cell is layout, not a question, so it must not become a slot.
    """
    if len(grid) < 2 or not any(not is_blank(text) for text in header):
        return False
    first_column = [row[0] for row in grid[1:] if row]
    if not first_column:
        return False
    labelled = sum(1 for cell in first_column if not is_blank(cell_text(cell)))
    return labelled * 2 >= len(first_column)


def _render_table(
    table: Table,
    namer: _SlotNamer,
    slots: list[Slot],
    max_cell_chars: int,
) -> list[str]:
    lines: list[str] = ["[TABLE]"]
    grid: list[list[_Cell]] = [_unique_cells(row) for row in table.rows]
    header = [cell_text(cell) for cell in grid[0]] if grid else []
    skip_corner = _has_corner_stub(grid, header)

    for row_index, cells in enumerate(grid):
        rendered: list[str] = []
        left_label = ""
        for column_index, cell in enumerate(cells):
            text = cell_text(cell)
            if is_blank(text) and row_index == 0 and column_index == 0 and skip_corner:
                rendered.append("")
                continue
            if not is_blank(text):
                rendered.append(text)
                left_label = text
                continue

            column_label = ""
            if row_index > 0 and column_index < len(header):
                column_label = header[column_index]
            slot = Slot(
                id=namer.next(),
                kind=KIND_CELL,
                paragraphs=(_writable_paragraph(cell),),
                label=left_label,
                column=column_label,
                max_chars=max_cell_chars,
            )
            slots.append(slot)
            rendered.append(MARKER_TEMPLATE.format(slot_id=slot.id))
        lines.append(" | ".join(rendered))

    lines.append("[/TABLE]")
    return lines


def extract(document: DocumentObject, max_cell_chars: int = 220) -> Extraction:
    """Scan `document` and return its marker text plus the ordered slot list."""
    namer = _SlotNamer()
    slots: list[Slot] = []
    lines: list[str] = []
    pending_blanks: list[Paragraph] = []

    def flush_blanks() -> None:
        """Turn a run of consecutive empty paragraphs into one answer slot."""
        if not pending_blanks:
            return
        slot = Slot(
            id=namer.next(),
            kind=KIND_PARA,
            paragraphs=tuple(pending_blanks),
        )
        slots.append(slot)
        lines.append(MARKER_TEMPLATE.format(slot_id=slot.id))
        pending_blanks.clear()

    for block in iter_block_items(document):
        if isinstance(block, Table):
            flush_blanks()
            lines.extend(_render_table(block, namer, slots, max_cell_chars))
            continue

        text = paragraph_text(block)
        if is_blank(text):
            pending_blanks.append(block)
            continue

        flush_blanks()
        if is_stub(text):
            slot = Slot(id=namer.next(), kind=KIND_STUB, paragraphs=(block,))
            slots.append(slot)
            lines.append(f"{text} {MARKER_TEMPLATE.format(slot_id=slot.id)}")
        else:
            lines.append(text)

    flush_blanks()
    return Extraction(text="\n".join(lines), slots=tuple(slots))


def strip_trailing_slots(extraction: Extraction, keep_last: int = 0) -> tuple[Slot, ...]:
    """Drop paragraph slots that trail the final piece of real content.

    Blank paragraphs at the very end of a document are layout padding, not
    questions, so they are not worth an answer.
    """
    slots = list(extraction.slots)
    while len(slots) > keep_last and slots[-1].kind == KIND_PARA:
        marker = MARKER_TEMPLATE.format(slot_id=slots[-1].id)
        if not extraction.text.rstrip().endswith(marker):
            break
        slots.pop()
    return tuple(slots)
