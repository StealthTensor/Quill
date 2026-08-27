"""Write generated answers back into the document, keeping its formatting.

Existing empty runs are reused wherever possible so the answer inherits the
font the worksheet author set for that blank. New paragraphs are cloned from
the blank they follow, so spacing and indentation survive.
"""

from __future__ import annotations

import copy

from docx.text.paragraph import Paragraph

from .slots import KIND_CELL, KIND_PARA, KIND_STUB, Slot


def _clear_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _add_run_like(paragraph: Paragraph, donor: Paragraph | None, text: str) -> None:
    """Add a run carrying `text`, copying character formatting from `donor`."""
    run = paragraph.add_run(text)
    donor_runs = donor.runs if donor is not None else []
    if donor_runs and donor_runs[0]._element.rPr is not None:
        run._element.insert(0, copy.deepcopy(donor_runs[0]._element.rPr))


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace a paragraph's content with `text`, keeping the first run's style."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run._element.getparent().remove(run._element)
        return
    paragraph.add_run(text)


def _append_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """Insert an empty clone of `paragraph` directly after it."""
    clone = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(clone)
    new_paragraph = Paragraph(clone, paragraph._parent)
    _clear_runs(new_paragraph)
    return new_paragraph


def _answer_lines(text: str) -> list[str]:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    return [line for line in lines if line] or [""]


def _write_paragraph_slot(slot: Slot, text: str) -> None:
    lines = _answer_lines(text)
    targets = list(slot.paragraphs)

    for index, line in enumerate(lines):
        if index < len(targets):
            _set_paragraph_text(targets[index], line)
            continue
        anchor = targets[-1]
        new_paragraph = _append_paragraph_after(anchor)
        _add_run_like(new_paragraph, slot.paragraphs[0], line)
        targets.append(new_paragraph)


def _write_stub_slot(slot: Slot, text: str) -> None:
    paragraph = slot.paragraphs[0]
    existing = paragraph.text
    separator = "" if existing.endswith((" ", "\t")) else " "
    _add_run_like(paragraph, paragraph, f"{separator}{text}")


def _write_cell_slot(slot: Slot, text: str) -> None:
    flattened = " ".join(part.strip() for part in text.split("\n") if part.strip())
    _set_paragraph_text(slot.paragraphs[0], flattened)


def apply_answers(slots: tuple[Slot, ...], answers: dict[str, str]) -> int:
    """Write every answered slot into the document. Returns how many were written."""
    written = 0
    for slot in slots:
        text = (answers.get(slot.id) or "").strip()
        if not text:
            continue

        if slot.kind == KIND_CELL:
            _write_cell_slot(slot, text)
        elif slot.kind == KIND_STUB:
            _write_stub_slot(slot, text)
        elif slot.kind == KIND_PARA:
            _write_paragraph_slot(slot, text)
        else:
            continue
        written += 1

    return written
